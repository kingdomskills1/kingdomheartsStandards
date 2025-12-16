import tkinter as tk
import re
from tkinter import filedialog, colorchooser, messagebox, ttk
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX
from tkinter import filedialog
import os


WORD_HIGHLIGHT_COLORS = {
    "None": None,                # No highlight
    "Yellow": WD_COLOR_INDEX.YELLOW,
    "Bright Green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "Red": WD_COLOR_INDEX.RED,
    "Blue": WD_COLOR_INDEX.BLUE,
    "Pink": WD_COLOR_INDEX.PINK,
    "Turquoise": WD_COLOR_INDEX.TURQUOISE,
    "Gray": WD_COLOR_INDEX.GRAY_25,
    "Dark Blue": WD_COLOR_INDEX.DARK_BLUE,
    "Dark Red": WD_COLOR_INDEX.DARK_RED,
    "Teal": WD_COLOR_INDEX.TEAL,
    "Violet": WD_COLOR_INDEX.VIOLET
}

# Color preview for GUI
PREVIEW_COLORS = {
    "None": "#ffffff",        # white for no highlight
    "Yellow": "#ffff00",
    "Bright Green": "#00ff00",
    "Red": "#ff0000",
    "Blue": "#0000ff",
    "Pink": "#ffc0cb",
    "Turquoise": "#40e0d0",
    "Gray": "#c0c0c0",
    "Dark Blue": "#00008b",
    "Dark Red": "#8b0000",
    "Teal": "#008080",
    "Violet": "#8a2be2"
}



def browse_target():
    if file_or_folder.get() == "file":
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    else:
        path = filedialog.askdirectory()
    file_path.set(path)


def paragraph_matches_filter(paragraph):
    filter_text = text_filter.get().strip()
    option = text_filter_option.get()
    use_regex = enable_regex.get()
    
    if not filter_text:
        return True  # no filter applied

    match_found = False
    if use_regex:
        try:
            # re.IGNORECASE makes it case-insensitive
            pattern = re.compile(filter_text, re.IGNORECASE)
            match_found = bool(pattern.search(paragraph.text))
        except re.error:
            match_found = False  # invalid regex -> no match
    else:
        match_found = filter_text.lower() in paragraph.text.lower()  # exact text match, case-insensitive

    if option == "Included":
        return match_found
    elif option == "Excluded":
        return not match_found



# Function to apply styles
def apply_text_style(file_path, font_name, font_size, bold, italic, underline, color,
                     include_headings, include_images, include_tables, option_choice, highlight):
    
    doc = Document(file_path)

    # ---------------- Run Styling Function ----------------
    def style_run(run):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.underline = WD_UNDERLINE.SINGLE if underline else None
        run.font.color.rgb = RGBColor(*color)
        # Apply highlight (None removes any previous highlight)
        run.font.highlight_color = highlight

    # ---------------- Check if paragraph has an image ----------------
    def paragraph_has_image(paragraph):
        return bool(paragraph._element.xpath('.//pic:pic'))

    # ---------------- Text Filter ----------------
    def paragraph_matches_filter(paragraph):
        filter_text_value = text_filter.get().strip()  # text from GUI entry
        filter_option = text_filter_option.get()       # Included / Excluded
        use_regex = enable_regex.get()                 # regex checkbox

        if not filter_text_value:
            return True  # no filter, match all

        match_found = False
        if use_regex:
            try:
                pattern = re.compile(filter_text_value, re.IGNORECASE)
                match_found = bool(pattern.search(paragraph.text))
            except re.error:
                match_found = False
        else:
            match_found = filter_text_value.lower() in paragraph.text.lower()

        if filter_option == "Included":
            return match_found
        elif filter_option == "Excluded":
            return not match_found

    # ---------------- Apply Styles Based on Option ----------------
    if option_choice == "Headings Only":
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading") and paragraph_matches_filter(paragraph):
                for run in paragraph.runs:
                    style_run(run)

    elif option_choice == "Tables Only":
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph_matches_filter(paragraph):
                            for run in paragraph.runs:
                                style_run(run)

    elif option_choice == "Images Only":
        for paragraph in doc.paragraphs:
            if paragraph_has_image(paragraph) and paragraph_matches_filter(paragraph):
                for run in paragraph.runs:
                    style_run(run)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph_has_image(paragraph) and paragraph_matches_filter(paragraph):
                            for run in paragraph.runs:
                                style_run(run)

    elif option_choice == "Text Only":
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                continue
            if paragraph_has_image(paragraph):
                continue
            in_table = any(paragraph in cell.paragraphs for table in doc.tables
                           for row in table.rows for cell in row.cells)
            if in_table:
                continue
            if paragraph_matches_filter(paragraph):
                for run in paragraph.runs:
                    style_run(run)

    elif option_choice == "All":
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading") and not include_headings.get():
                continue
            if paragraph_has_image(paragraph) and not include_images.get():
                continue
            if paragraph_matches_filter(paragraph):
                for run in paragraph.runs:
                    style_run(run)
        if include_tables.get():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.style.name.startswith("Heading") and not include_headings.get():
                                continue
                            if paragraph_has_image(paragraph) and not include_images.get():
                                continue
                            if paragraph_matches_filter(paragraph):
                                for run in paragraph.runs:
                                    style_run(run)

    # ---------------- Save ----------------
    doc.save(file_path)
    # messagebox.showinfo("Success", f"Styles applied to:\n{file_path}")


def on_highlight_select(event=None):
    global highlight_color
    global PREVIEW_COLORS
    selected = highlight_var.get()
    highlight_color = WORD_HIGHLIGHT_COLORS[selected]
    highlight_color_preview.config(bg=PREVIEW_COLORS[selected])





    # Preview exact Word color
    PREVIEW_COLORS = {
        "Yellow": "#ffff00",
        "Bright Green": "#00ff00",
        "Red": "#ff0000",
        "Blue": "#0000ff",
        "Pink": "#ffc0cb",
        "Turquoise": "#40e0d0",
        "Gray": "#c0c0c0",
        "Dark Blue": "#00008b",
        "Dark Red": "#8b0000",
        "Teal": "#008080",
        "Violet": "#8a2be2"
    }

    highlight_color_preview.config(bg=PREVIEW_COLORS[selected])


# GUI Setup
root = tk.Tk()
root.title("Word Text Styler")
root.geometry("480x700")

file_path = tk.StringVar()
font_name = tk.StringVar(value="Calibri")
font_size = tk.IntVar(value=12)
bold = tk.BooleanVar()
italic = tk.BooleanVar()
underline = tk.BooleanVar()
include_headings = tk.BooleanVar(value=True)
include_images = tk.BooleanVar(value=True)
include_tables = tk.BooleanVar(value=True)
text_color = (0, 0, 0)
option_var = tk.StringVar(value="All")

highlight_color = None
highlight_enabled = tk.BooleanVar(value=False)

highlight_enabled = tk.BooleanVar(value=False)
highlight_var = tk.StringVar(value="Yellow")
highlight_color = WD_COLOR_INDEX.YELLOW

text_filter = tk.StringVar()
text_filter_option = tk.StringVar(value="Included")

enable_regex = tk.BooleanVar(value=False)  # default unchecked

file_or_folder = tk.StringVar(value="file")  # default to file


def update_checkboxes(*args):
    choice = option_var.get()
    if choice == "All":
        cb_headings.config(state="normal")
        cb_images.config(state="normal")
        cb_tables.config(state="normal")
    else:
        cb_headings.config(state="disabled")
        cb_images.config(state="disabled")
        cb_tables.config(state="disabled")

# GUI Functions
def browse_file():
    path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    file_path.set(path)

def choose_color():
    global text_color
    color_code = colorchooser.askcolor(title="Choose Text Color")
    if color_code[0]:
        text_color = tuple(int(c) for c in color_code[0])
        text_color_preview.config(bg=color_code[1])


def choose_highlight():
    global highlight_color
    color_code = colorchooser.askcolor(title="Choose Highlight Color")
    if color_code[0]:
        # Map approximate RGB to nearest WD_COLOR_INDEX
        # Limited predefined colors
        r, g, b = [int(c) for c in color_code[0]]
        # Simplified mapping for demo purposes
        highlight_color_preview.config(bg=color_code[1])
        highlight_color = WD_COLOR_INDEX.YELLOW  # Can be expanded for more colors
        

def apply_styles():
    path = file_path.get()
    if not path:
        messagebox.showerror("Error", "Please select a file or folder")
        return

    docx_files = []

    if file_or_folder.get() == "file":
        docx_files = [path]
    else:  # folder
        if include_subfolders.get():
            for root_dir, dirs, files in os.walk(path):
                for f in files:
                    if f.lower().endswith(".docx"):
                        docx_files.append(os.path.join(root_dir, f))
        else:
            for f in os.listdir(path):
                full_path = os.path.join(path, f)
                if os.path.isfile(full_path) and f.lower().endswith(".docx"):
                    docx_files.append(full_path)

    if not docx_files:
        messagebox.showinfo("Info", "No .docx files found.")
        return

    # Apply styles to all files
    for file in docx_files:
        apply_text_style(
            file,
            font_name.get(),
            font_size.get(),
            bold.get(),
            italic.get(),
            underline.get(),
            text_color,
            include_headings,
            include_images,
            include_tables,
            option_var.get(),
            highlight_color
        )

    # Show ONE success message after processing
    if file_or_folder.get() == "file":
        messagebox.showinfo("Success", f"Styles applied to:\n{docx_files[0]}")
    else:
        messagebox.showinfo(
            "Success",
            f"Styles applied to {len(docx_files)} file(s) in the folder:\n{path}"
        )



# ================= GUI WIDGETS =================

# ================= File/Folder Selection =================
tk.Label(root, text="Select Target:").pack(pady=5)

# Radio buttons to select File or Folder
file_or_folder = tk.StringVar(value="file")  # default to file

tk.Radiobutton(root, text="File", variable=file_or_folder, value="file", command=lambda: update_subfolder_state()).pack()
tk.Radiobutton(root, text="Folder", variable=file_or_folder, value="folder", command=lambda: update_subfolder_state()).pack()

# Entry and Browse button
tk.Entry(root, textvariable=file_path, width=45).pack(pady=5)
tk.Button(root, text="Browse", command=lambda: browse_target()).pack(pady=5)

# Subfolders checkbox
include_subfolders = tk.BooleanVar(value=False)
cb_subfolders = tk.Checkbutton(root, text="Include Subfolders", variable=include_subfolders)
cb_subfolders.pack(pady=5)

# Function to enable/disable subfolders checkbox
def update_subfolder_state():
    if file_or_folder.get() == "file":
        cb_subfolders.config(state="disabled")
    else:
        cb_subfolders.config(state="normal")

update_subfolder_state()  # initialize checkbox state

# ================= Font Selection =================
tk.Label(root, text="Font Name:").pack()
common_fonts = [
    "Arial", "Calibri", "Times New Roman", "Verdana", "Tahoma",
    "Courier New", "Georgia", "Trebuchet MS", "Impact", "Comic Sans MS"
]
font_combobox = ttk.Combobox(root, textvariable=font_name, values=common_fonts)
font_combobox.pack(pady=5)
font_combobox['state'] = 'normal'

tk.Label(root, text="Font Size:").pack()
common_sizes = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36, 48, 72]
size_combobox = ttk.Combobox(root, textvariable=font_size, values=common_sizes)
size_combobox.pack(pady=5)
size_combobox['state'] = 'normal'

# ================= Formatting =================
tk.Checkbutton(root, text="Bold", variable=bold).pack()
tk.Checkbutton(root, text="Italic", variable=italic).pack()
tk.Checkbutton(root, text="Underline", variable=underline).pack()

# ================= Apply Options =================
tk.Label(root, text="Apply Options:").pack(pady=5)
option_menu = ttk.Combobox(
    root,
    textvariable=option_var,
    values=["All", "Headings Only", "Tables Only", "Images Only", "Text Only"]
)
option_menu.pack(pady=5)
option_menu.bind("<<ComboboxSelected>>", update_checkboxes)

# Checkboxes
cb_headings = tk.Checkbutton(root, text="Include Headings", variable=include_headings)
cb_headings.pack(pady=2)

cb_images = tk.Checkbutton(root, text="Include Images", variable=include_images)
cb_images.pack(pady=2)

cb_tables = tk.Checkbutton(root, text="Include Tables", variable=include_tables)
cb_tables.pack(pady=2)

# ================= Color Selection =================
tk.Button(root, text="Choose Text Color", command=choose_color).pack(pady=5)

# Preview frame
preview_frame = tk.Frame(root)
preview_frame.pack(pady=5)

tk.Label(preview_frame, text="Text Color:").grid(row=0, column=0, padx=5)
text_color_preview = tk.Label(preview_frame, width=4, height=1, bg="black", relief="solid")
text_color_preview.grid(row=0, column=1, padx=5)

tk.Label(preview_frame, text="Highlight:").grid(row=0, column=2, padx=5)
highlight_combo = ttk.Combobox(
    preview_frame,
    textvariable=highlight_var,
    values=list(WORD_HIGHLIGHT_COLORS.keys()),  # includes "None"
    state="readonly",
    width=12
)
highlight_combo.grid(row=0, column=3, padx=5)
highlight_combo.bind("<<ComboboxSelected>>", on_highlight_select)
highlight_combo.current(1)  # default "Yellow"

highlight_color_preview = tk.Label(preview_frame, width=4, height=1, bg="yellow", relief="solid")
highlight_color_preview.grid(row=0, column=4, padx=5)

# ================= Text Filter Section =================
tk.Label(root, text="Text Filter:").pack(pady=5)

text_filter_entry = tk.Entry(root, textvariable=text_filter, width=40)
text_filter_entry.pack(pady=5)

tk.Checkbutton(root, text="Enable Regex", variable=enable_regex).pack(pady=5)

text_filter_option_menu = ttk.Combobox(
    root,
    textvariable=text_filter_option,
    values=["Included", "Excluded"],
    state="readonly"
)
text_filter_option_menu.pack(pady=5)
text_filter_option_menu.current(0)  # default "Included"

# ================= Apply Button =================
tk.Button(root, text="Apply Styles", command=apply_styles).pack(pady=10)

# Initialize states
update_checkboxes()
on_highlight_select()



root.mainloop()

