from docx import Document
import os
from collections import Counter


# -----------------------------
# Insert in text or DOCX
# -----------------------------
def insert_at_matches(
    doc_or_text,
    matches_list,
    insert_type,
    content="",
    repeat=1,
    position="before",
    line_offset=0,
    insert_reference="matched_line",
    content_type="all",    
):
    insert_text = {"space": " ", "newline": "\n"}.get(insert_type.lower(), content)
    prefix = insert_text * repeat
    total_insertions = 0

    # ================= TXT =================
    if isinstance(doc_or_text, str):
        lines = doc_or_text.splitlines()
        new_lines = lines[:]

        for line_no, line_text, _ in matches_list:
            idx = line_no - 1 + line_offset
            if 0 <= idx < len(new_lines):
                new_lines[idx] = prefix + new_lines[idx]
                total_insertions += 1

        return total_insertions, "\n".join(new_lines)

    # ================= DOCX =================
    doc = doc_or_text
    if insert_reference == "matched_line":
        if position.lower() == "before":
            # ---------- HANDLE NORMAL PARAGRAPHS ----------
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                if len(p._element.xpath(".//w:drawing")) > 0:
                    continue

                # Check content type
                if content_type.lower() == "text":
                    # Only normal paragraphs (not headings)
                    if p.style and p.style.name.startswith("Heading"):
                        continue
                elif content_type.lower() in ("headings", "tables_headings"):
                    # Only heading paragraphs
                    if not (p.style and p.style.name.startswith("Heading")):
                        continue
                # else: "all" -> everything, no need to check

                # Count total occurrences for this paragraph
                match_count = 0
                for line_no, line_text, ranges in matches_list:
                    if line_text == p.text:
                        match_count += len(ranges)  # number of matches within the line

                if match_count == 0:
                    continue

                if p.runs:
                    p.runs[0].text = (insert_text * repeat * match_count) + p.runs[0].text
                    total_insertions += match_count

            if content_type.lower() in ("all", "tables", "tables_headings"):
                # ---------- HANDLE TABLES ----------
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if not p.text.strip():
                                    continue
                                if len(p._element.xpath(".//w:drawing")) > 0:
                                    continue

                                # Count total occurrences in this paragraph
                                match_count = 0
                                for line_no, line_text, ranges in matches_list:
                                    if line_text == p.text:
                                        match_count += len(ranges)

                                if match_count == 0:
                                    continue

                                if p.runs:
                                    p.runs[0].text = (insert_text * repeat * match_count) + p.runs[0].text
                                    total_insertions += match_count
        if position.lower() == "after":
            print("after")
            # ---------- HANDLE NORMAL PARAGRAPHS ----------
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                if len(p._element.xpath(".//w:drawing")) > 0:
                    continue

                # Check content type
                if content_type.lower() == "text":
                    if p.style and p.style.name.startswith("Heading"):
                        continue
                elif content_type.lower() in ("headings", "tables_headings"):
                    if not (p.style and p.style.name.startswith("Heading")):
                        continue

                match_count = 0
                for line_no, line_text, ranges in matches_list:
                    if line_text == p.text:
                        match_count += len(ranges)

                if match_count == 0:
                    continue

                if p.runs:
                    p.runs[-1].text = p.runs[-1].text + (insert_text * repeat * match_count)
                    total_insertions += match_count

            if content_type.lower() in ("all", "tables", "tables_headings"):
                # ---------- HANDLE TABLES ----------
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if not p.text.strip():
                                    continue
                                if len(p._element.xpath(".//w:drawing")) > 0:
                                    continue

                                match_count = 0
                                for line_no, line_text, ranges in matches_list:
                                    if line_text == p.text:
                                        match_count += len(ranges)

                                if match_count == 0:
                                    continue

                                if p.runs:
                                    p.runs[-1].text = p.runs[-1].text + (insert_text * repeat * match_count)
                                    total_insertions += match_count
    if insert_reference == "matched_content":
        if position.lower() == "before":
            # ---------- HANDLE NORMAL PARAGRAPHS ----------
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                if len(p._element.xpath(".//w:drawing")) > 0:
                    continue

                # Check content type
                if content_type.lower() == "text":
                    if p.style and p.style.name.startswith("Heading"):
                        continue
                elif content_type.lower() in ("headings", "tables_headings"):
                    if not (p.style and p.style.name.startswith("Heading")):
                        continue

                full_text = "".join(run.text for run in p.runs)
                match_count = 0

                for line_no, line_text, ranges in matches_list:
                    if line_text == full_text:
                        for start, end in sorted(ranges, reverse=True):
                            # find run containing start
                            char_pos = 0
                            for run in p.runs:
                                run_len = len(run.text)
                                if char_pos <= start < char_pos + run_len:
                                    offset = start - char_pos
                                    run.text = (
                                        run.text[:offset]
                                        + (insert_text * repeat)
                                        + run.text[offset:]
                                    )
                                    match_count += 1
                                    break
                                char_pos += run_len

                if match_count > 0:
                    total_insertions += match_count

            if content_type.lower() in ("all", "tables", "tables_headings"):
                # ---------- HANDLE TABLES ----------
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if not p.text.strip():
                                    continue
                                if len(p._element.xpath(".//w:drawing")) > 0:
                                    continue

                                full_text = "".join(run.text for run in p.runs)
                                match_count = 0

                                for line_no, line_text, ranges in matches_list:
                                    if line_text == full_text:
                                        for start, end in sorted(ranges, reverse=True):
                                            char_pos = 0
                                            for run in p.runs:
                                                run_len = len(run.text)
                                                if char_pos <= start < char_pos + run_len:
                                                    offset = start - char_pos
                                                    run.text = (
                                                        run.text[:offset]
                                                        + (insert_text * repeat)
                                                        + run.text[offset:]
                                                    )
                                                    match_count += 1
                                                    break
                                                char_pos += run_len

                                if match_count > 0:
                                    total_insertions += match_count
        if position.lower() == "after":
            # ---------- HANDLE NORMAL PARAGRAPHS ----------
            for p in doc.paragraphs:
                if not p.text.strip():
                    continue
                if len(p._element.xpath(".//w:drawing")) > 0:
                    continue

                # Check content type
                if content_type.lower() == "text":
                    if p.style and p.style.name.startswith("Heading"):
                        continue
                elif content_type.lower() in ("headings", "tables_headings"):
                    if not (p.style and p.style.name.startswith("Heading")):
                        continue

                full_text = "".join(run.text for run in p.runs)
                match_count = 0

                for line_no, line_text, ranges in matches_list:
                    if line_text == full_text:
                        for start, end in sorted(ranges, reverse=True):
                            char_pos = 0
                            for run in p.runs:
                                run_len = len(run.text)
                                if char_pos <= end <= char_pos + run_len:
                                    offset = end - char_pos
                                    run.text = (
                                        run.text[:offset]
                                        + (insert_text * repeat)
                                        + run.text[offset:]
                                    )
                                    match_count += 1
                                    break
                                char_pos += run_len

                if match_count > 0:
                    total_insertions += match_count

            if content_type.lower() in ("all", "tables", "tables_headings"):
                # ---------- HANDLE TABLES ----------
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if not p.text.strip():
                                    continue
                                if len(p._element.xpath(".//w:drawing")) > 0:
                                    continue

                                full_text = "".join(run.text for run in p.runs)
                                match_count = 0

                                for line_no, line_text, ranges in matches_list:
                                    if line_text == full_text:
                                        for start, end in sorted(ranges, reverse=True):
                                            char_pos = 0
                                            for run in p.runs:
                                                run_len = len(run.text)
                                                if char_pos <= end <= char_pos + run_len:
                                                    offset = end - char_pos
                                                    run.text = (
                                                        run.text[:offset]
                                                        + (insert_text * repeat)
                                                        + run.text[offset:]
                                                    )
                                                    match_count += 1
                                                    break
                                                char_pos += run_len

                                if match_count > 0:
                                    total_insertions += match_count



    return total_insertions

# -----------------------------
# Apply insert to all files
# -----------------------------
def insert_all_files(gui, search_path_func):
    locked_files = []
    total_insertions = 0
    per_file_insertions = []

    results = search_path_func(
        path=gui.entry_path.get(),
        search_text=gui.entry_search_text.get(),
        selected_type=gui.selected_type.get(),
        case_sensitive=gui.case_sensitive_var.get(),
        use_regex=gui.enable_regex_var.get(),
        search_subfolders=gui.subfolders_var.get(),
        txt_only=gui.txt_var.get(),
        doc_only=gui.doc_var.get(),
        pdf_only=gui.pdf_var.get(),
        content_type=gui.content_type_var.get()
    )

    if not results:
        return 0, []

    insert_type = gui.insert_content_type_var.get()
    insert_reference = gui.insert_reference_var.get()
    content = gui.insert_content_text_var.get()
    try:
        repeat = int(gui.insert_repeat_var.get())
    except:
        repeat = 1
    position = gui.insert_position_var.get()
    try:
        line_offset = int(gui.insert_line_offset_var.get())
    except:
        line_offset = 0

    for result in results:
        if result.get("unsupported") or result.get("error"):
            continue
        matches = result.get("matches", [])
        if not matches:
            continue

        file_path = result["file_path"]
        ext = os.path.splitext(file_path)[1].lower()
        count = 0

        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                count, new_text = insert_at_matches(
                    doc_or_text=text,
                    matches_list=matches,
                    insert_type=insert_type,
                    content=content,
                    repeat=repeat,
                    position=position,
                    line_offset=line_offset,
                    insert_reference=insert_reference,
                    content_type=gui.content_type_var.get()
                )
                if count > 0:
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(new_text)

            elif ext == ".docx":
                doc = Document(file_path)
                count = insert_at_matches(
                    doc_or_text=doc,
                    matches_list=matches,
                    insert_type=insert_type,
                    content=content,
                    repeat=repeat,
                    position=position,
                    line_offset=line_offset,
                    insert_reference=insert_reference,
                    content_type=gui.content_type_var.get()
                )
                if count > 0:
                    doc.save(file_path)
        except PermissionError:
            # File is open in another program
            locked_files.append(file_path)
            continue

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            continue


        if count > 0:
            total_insertions += count
            per_file_insertions.append((file_path, count))
    
    return total_insertions, per_file_insertions, locked_files
