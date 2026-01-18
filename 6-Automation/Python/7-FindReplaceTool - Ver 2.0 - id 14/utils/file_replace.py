import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import tkinter.font as tkFont
import unicodedata
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----------------- Helpers -----------------
def extract_plain_text(tk_text_widget):
    """
    Extract plain text from the Replace Text widget.
    """
    return tk_text_widget.get("1.0", "end-1c")


def normalize_text(s):
    """
    Normalize Word text to handle special characters like:
    - non-breaking hyphens
    - non-breaking spaces
    - Unicode normalization
    """
    s = s.replace("\u2011", "-").replace("\u00A0", " ")
    s = unicodedata.normalize("NFKC", s)
    return s


def get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf):
    """
    Return a list of files to process.
    """
    files = []
    if is_file:
        return [path]

    allowed_exts = []
    if txt: allowed_exts.append(".txt")
    if doc: allowed_exts.append(".docx")
    if pdf: allowed_exts.append(".pdf")

    for root, dirs, filenames in os.walk(path):
        for name in filenames:
            if not allowed_exts or any(name.lower().endswith(ext) for ext in allowed_exts):
                files.append(os.path.join(root, name))
        if not include_subfolders:
            break

    return files


# ----------------- TXT Replacement -----------------
def replace_in_file_txt(path, find_text, replace_text, case_sensitive, regex):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    flags = 0 if case_sensitive else re.IGNORECASE

    if regex:
        new_content = re.sub(find_text, replace_text, content, flags=flags)
    else:
        pattern = re.compile(re.escape(find_text), flags=flags)
        new_content = pattern.sub(replace_text, content)

    with open(path, "w", encoding="utf-8", errors="ignore") as f:
        f.write(new_content)



HIGHLIGHT_MAP = {
    0: "None", 1: "Black", 2: "Blue", 3: "Turquoise",
    4: "Bright Green", 5: "Pink", 6: "Red", 7: "Yellow", 8: "White",
    9: "Dark Blue", 10: "Teal", 11: "Green", 12: "Violet",
    13: "Dark Red", 14: "Dark Yellow", 15: "Gray50", 16: "Gray25"
}

WORD_TO_COLOR = {
    "Black": WD_COLOR_INDEX.BLACK, "Blue": WD_COLOR_INDEX.BLUE,
    "Turquoise": WD_COLOR_INDEX.TURQUOISE, "Bright Green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "Pink": WD_COLOR_INDEX.PINK, "Red": WD_COLOR_INDEX.RED, "Yellow": WD_COLOR_INDEX.YELLOW,
    "White": WD_COLOR_INDEX.WHITE, "Dark Blue": WD_COLOR_INDEX.DARK_BLUE, 
    "Teal": WD_COLOR_INDEX.TEAL, "Green": WD_COLOR_INDEX.GREEN, 
    "Violet": WD_COLOR_INDEX.VIOLET, "Dark Red": WD_COLOR_INDEX.DARK_RED,
    "Dark Yellow": WD_COLOR_INDEX.DARK_YELLOW, "Gray50": WD_COLOR_INDEX.GRAY_50, 
    "Gray25": WD_COLOR_INDEX.GRAY_25, "None": None
}

# Map Tkinter background colors to DOCX highlight colors if needed
BG_COLOR_TO_DOCX_HIGHLIGHT = {
    "#000000": WD_COLOR_INDEX.BLACK,
    "#0000FF": WD_COLOR_INDEX.BLUE,
    "#00FFFF": WD_COLOR_INDEX.TURQUOISE,
    "#00FF00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#FFC0CB": WD_COLOR_INDEX.PINK,
    "#FF0000": WD_COLOR_INDEX.RED,
    "#FFFF00": WD_COLOR_INDEX.YELLOW,
    "#FFFFFF": WD_COLOR_INDEX.WHITE,
    "#00008B": WD_COLOR_INDEX.DARK_BLUE,
    "#008080": WD_COLOR_INDEX.TEAL,
    "#008000": WD_COLOR_INDEX.GREEN,
    "#EE82EE": WD_COLOR_INDEX.VIOLET,
    "#8B0000": WD_COLOR_INDEX.DARK_RED,
    "#9B870C": WD_COLOR_INDEX.DARK_YELLOW,
    "#808080": WD_COLOR_INDEX.GRAY_50,
    "#C0C0C0": WD_COLOR_INDEX.GRAY_25,
}


WORD_HEADING_FORMATS = {
    "Heading 1": {"size": 20, "bold": False},
    "Heading 2": {"size": 16, "bold": False},
    "Heading 3": {"size": 14, "bold": False},
    "Heading 4": {"size": 12, "bold": False},
    "Heading 5": {"size": 12, "bold": False, "italic": True},
    "Heading 6": {"size": 12, "bold": False, "italic": True},
    "Heading 7": {"size": 12, "bold": False, "italic": True},
    "Heading 8": {"size": 12, "bold": False, "italic": True},
    "Heading 9": {"size": 12, "bold": False, "italic": True},
}

WORD_HEADING_BLUE = RGBColor(0xB3, 0xC1, 0xD9)  # #B3C1D9

def apply_word_heading_format(para):
    """
    Apply Word default Heading formatting (Heading 1–9) exactly as Word does.
    - Sets paragraph style to Heading style.
    - Clears any run-level formatting (bold, italic, size, color, font).
    """
    if not para.style:
        return

    style_name = para.style.name
    if style_name not in WORD_HEADING_FORMATS:
        return

    fmt = WORD_HEADING_FORMATS[style_name]

    # --- Apply paragraph style itself ---
    doc_styles = para.part.styles
    if style_name in doc_styles:
        para.style = doc_styles[style_name]

    # --- Clear run-level formatting so Word defaults are applied ---
    for run in para.runs:
        run.font.bold = None
        run.font.italic = None
        run.font.size = None
        run.font.color.rgb = None
        run.font.color.theme_color = None
        run.font.name = None


def build_plain_char_formats(text):
    """
    Build minimal char_formats list so replace_paragraph_safe_inplace
    inserts text instead of deleting it.
    Formatting will be overridden later.
    """
    return [{
        "text": ch,
        "font_family": "Arial",
        "font_size": 12,
        "bold": False,
        "italic": False,
        "color": None,
        "highlight": None,
        "highlight_hex": None
    } for ch in text]


def set_run_shading(run, hex_color):
    """
    Apply shading to a run using a hex color (e.g., #C0C0C0)
    """
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        rPr.append(shd)
    shd.set(qn('w:fill'), hex_color.replace('#',''))

def bg_to_docx_highlight(bg: str):
    """
    Convert Tkinter background hex color to DOCX highlight.
    Handles case-insensitive matching and fallback for common grey shades.
    """
    if not bg:
        return None
    bg = bg.upper()
    mapping = {
        "#000000": WD_COLOR_INDEX.BLACK,
        "#0000FF": WD_COLOR_INDEX.BLUE,
        "#00FFFF": WD_COLOR_INDEX.TURQUOISE,
        "#00FF00": WD_COLOR_INDEX.BRIGHT_GREEN,
        "#FFC0CB": WD_COLOR_INDEX.PINK,
        "#FF0000": WD_COLOR_INDEX.RED,
        "#FFFF00": WD_COLOR_INDEX.YELLOW,
        "#FFFFFF": WD_COLOR_INDEX.WHITE,
        "#00008B": WD_COLOR_INDEX.DARK_BLUE,
        "#008080": WD_COLOR_INDEX.TEAL,
        "#008000": WD_COLOR_INDEX.GREEN,
        "#EE82EE": WD_COLOR_INDEX.VIOLET,
        "#8B0000": WD_COLOR_INDEX.DARK_RED,
        "#9B870C": WD_COLOR_INDEX.DARK_YELLOW,
        "#808080": WD_COLOR_INDEX.GRAY_50,
        "#C0C0C0": WD_COLOR_INDEX.GRAY_25,
    }
    if bg in mapping:
        return mapping[bg]
    # fallback for other common grey shades
    if bg in ("#BEBEBE", "#D3D3D3"):
        return WD_COLOR_INDEX.GRAY_25
    return None


def get_text_widget_char_formats(text_widget):
    """
    Extract per-character text and formatting from a Tkinter Text widget.
    Returns a list of dictionaries for use in replace_paragraph_safe_inplace.
    """
    result = []
    text = text_widget.get("1.0", "end-1c")
    for i, ch in enumerate(text):
        idx = f"1.0 + {i} chars"
        # Default formatting
        font_family = "Arial"
        font_size = 12
        bold = False
        italic = False
        fg_color = "#000000"
        bg_color = None
        highlight_hex = None

        # Check all tags at this position
        tags = text_widget.tag_names(idx)
        for tag in tags:
            tag_config = text_widget.tag_cget(tag, "font")
            if tag_config:
                tk_font = tkFont.Font(font=text_widget.tag_cget(tag, "font"))
                font_family = tk_font.actual("family")
                font_size = tk_font.actual("size")
                bold = tk_font.actual("weight") == "bold"
                italic = tk_font.actual("slant") == "italic"

            fg = text_widget.tag_cget(tag, "foreground")
            if fg:
                fg_color = fg

            bg = text_widget.tag_cget(tag, "background")
            if bg:
                bg_color = bg_to_docx_highlight(bg)
                highlight_hex = bg  # store original hex for fallback

        # Convert fg_color #RRGGBB to RGBColor
        r = int(fg_color[1:3], 16)
        g = int(fg_color[3:5], 16)
        b = int(fg_color[5:7], 16)
        color_rgb = RGBColor(r, g, b)

        result.append({
            "text": ch,
            "font_family": font_family,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "color": color_rgb,
            "highlight": bg_color,
            "highlight_hex": highlight_hex
        })
    return result

# The function replace_paragraph_safe_inplace is responsible for replacing text within a Word paragraph (para) in-place 
# while preserving formatting and highlights.
def replace_paragraph_safe_inplace(para, pattern, char_formats):
    """
    Replaces only matched text in-place with per-character formatting.
    If char_formats is empty, removes the matched text.
    Unmatched text keeps all original formatting, including highlight and color.
    """
    if not para.text:
        return 0

    full_text = para.text
    matches = list(pattern.finditer(full_text))
    if not matches:
        return  0 # No matches, keep as is

    def force_run_font(run, font_name):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for k in ("ascii", "hAnsi", "eastAsia", "cs"):
            rFonts.set(qn(f"w:{k}"), font_name)
        run.font.name = font_name

    # Build mapping: run -> paragraph text positions
    runs = []
    pos = 0
    for run in para.runs:
        runs.append({
            "run": run,
            "start": pos,
            "end": pos + len(run.text),
            "font_name": run.font.name,
            "size": run.font.size,
            "bold": run.font.bold,
            "italic": run.font.italic,
            "color": run.font.color.rgb if run.font.color else None,
            "highlight": run.font.highlight_color
        })
        pos += len(run.text)

    new_runs = []
    cursor = 0

    for match in matches:
        start, end = match.span()

        # Text before match -> original formatting
        for r in runs:
            if r["end"] <= cursor or r["start"] >= start:
                continue
            s = max(cursor, r["start"])
            e = min(start, r["end"])
            if s < e:
                frag = r["run"].text[s - r["start"]: e - r["start"]]
                new_runs.append((frag, r, False))

        # Matched replacement
        if char_formats:
            for c in char_formats:
                new_runs.append((c["text"], c, True))

        cursor = end

    # Text after last match -> original formatting
    for r in runs:
        if r["end"] <= cursor:
            continue
        frag = r["run"].text[max(cursor - r["start"], 0):]
        if frag:
            new_runs.append((frag, r, False))

    # Clear paragraph and rebuild runs
    para.clear()
    for text, fmt, is_match in new_runs:
        r = para.add_run(text)
        if is_match:
            r.font.name = fmt["font_family"]
            r.font.size = Pt(fmt["font_size"])
            r.font.bold = fmt["bold"]
            r.font.italic = fmt["italic"]
            if fmt["color"]:
                r.font.color.rgb = fmt["color"]
            # Handle highlight
            if fmt.get("highlight") is not None:
                try:
                    r.font.highlight_color = fmt["highlight"]
                except Exception:
                    if fmt.get("highlight_hex"):
                        set_run_shading(r, fmt["highlight_hex"])
            elif fmt.get("highlight_hex"):
                # If highlight not in WD_COLOR_INDEX, still apply shading
                set_run_shading(r, fmt["highlight_hex"])
            force_run_font(r, fmt["font_family"])
        else:
            r.font.name = fmt["font_name"]
            r.font.size = fmt["size"]
            r.font.bold = fmt["bold"]
            r.font.italic = fmt["italic"]
            if fmt["color"]:
                r.font.color.rgb = fmt["color"]
            if fmt.get("highlight") is not None:
                try:
                    r.font.highlight_color = fmt["highlight"]
                except Exception:
                    if fmt.get("highlight_hex"):
                        set_run_shading(r, fmt["highlight_hex"])
            elif fmt.get("highlight_hex"):
                set_run_shading(r, fmt["highlight_hex"])

    return len(matches)


def process_tables(tables, handler):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    handler(para)
                if cell.tables:
                    process_tables(cell.tables, handler)

def replace_in_file_docx(path, tk_replace_widget, find_text,
                         case_sensitive=False, regex=False):
    """
    Replaces matched text in DOCX with Word selection formatting.
    If no text selected, removes the matched text.
    Only matched text is replaced; unmatched text keeps original formatting including highlight.
    """
    doc = Document(path)
    char_formats = get_text_widget_char_formats(tk_replace_widget)
    if char_formats is None:
        char_formats = []  # Treat as empty to remove matched text

    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(find_text, flags) if regex else re.compile(re.escape(find_text), flags)

    handler = lambda p: replace_paragraph_safe_inplace(p, pattern, char_formats)

    for para in doc.paragraphs:
        handler(para)
    process_tables(doc.tables, handler)

    doc.save(path)
    print(f"Replaced text in {path}.")
# end functions of replaced

# ----------------- Main Replace Process -----------------
def run_replace_process(gui, path, is_file, find_text, tk_replace_widget,
                        case_sensitive, regex, include_subfolders, txt, doc, pdf, content_type):

    gui.text_results.config(state="normal")
    
    # Ensure text_results is cleared at the beginning (optional based on your flow)
    # gui.text_results.delete("1.0", "end")  # You might not want to clear, depending on your need

    total_replacements = 0
    files_with_replacements = 0

    files = get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf)
    gui.text_results.insert("end", f"Processing {len(files)} file(s)...\n\n")

    replace_plain_text = extract_plain_text(tk_replace_widget)  # Extract replacement text

    for file_path in files:
        file_replacements = 0

        try:
            # ---------- TXT ---------- 
            if file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()

                flags = 0 if case_sensitive else re.IGNORECASE

                if regex:
                    new_content, count = re.subn(find_text, replace_plain_text, content, flags=flags)
                else:
                    pattern = re.compile(re.escape(find_text), flags=flags)
                    new_content, count = pattern.subn(replace_plain_text, content)

                with open(file_path, "w", encoding="utf-8", errors="ignore") as f:
                    f.write(new_content)

                file_replacements = count


            # ---------- DOCX ----------
            elif file_path.lower().endswith(".docx"):
                docx = Document(file_path)
                char_formats = get_text_widget_char_formats(tk_replace_widget) or []

                flags = 0 if case_sensitive else re.IGNORECASE
                pattern = re.compile(find_text, flags) if regex else re.compile(re.escape(find_text), flags)

                # Extract replacement text
                replace_plain_text = extract_plain_text(tk_replace_widget)
                apply_heading_format = gui.apply_heading_format_var.get()

                # --- Process each paragraph ---
                for para in docx.paragraphs:
                    is_heading = para.style and para.style.name.startswith("Heading")

                    if is_heading:
                        if apply_heading_format:
                            # Build simple char_formats list to avoid deleting text
                            plain_formats = build_plain_char_formats(replace_plain_text)
                            file_replacements += replace_paragraph_safe_inplace(para, pattern, plain_formats)
                            # Apply Word default Heading format
                            apply_word_heading_format(para)
                        else:
                            if content_type in ("all", "headings", "tables_headings"):
                                file_replacements += replace_paragraph_safe_inplace(para, pattern, char_formats)
                    else:
                        # Non-heading text
                        if content_type in ("all", "text"):
                            file_replacements += replace_paragraph_safe_inplace(para, pattern, char_formats)

                # --- Process tables ---
                if content_type in ("tables", "tables_headings", "all"):
                    def table_handler(p):
                        nonlocal file_replacements
                        is_heading = p.style and p.style.name.startswith("Heading")
                        if is_heading:
                            if apply_heading_format:
                                plain_formats = build_plain_char_formats(replace_plain_text)
                                file_replacements += replace_paragraph_safe_inplace(p, pattern, plain_formats)
                                apply_word_heading_format(p)
                            else:
                                file_replacements += replace_paragraph_safe_inplace(p, pattern, char_formats)
                        else:
                            file_replacements += replace_paragraph_safe_inplace(p, pattern, char_formats)
                    process_tables(docx.tables, table_handler)

                docx.save(file_path)

            # ---------- PDF ---------- 
            elif file_path.lower().endswith(".pdf"):
                gui.text_results.insert("end", f"PDF replace not supported: {file_path}\n")
                continue

            total_replacements += file_replacements
            if file_replacements > 0:
                files_with_replacements += 1

            # Append file replacement results and add a new line after each result
            gui.text_results.insert(
                "end",
                f"{file_path} ({file_replacements} replacements)\n"
            )

        except Exception as e:
            gui.text_results.insert("end", f"Error in {file_path}: {e}\n")

    # ---------- SUMMARY ----------
    gui.text_results.insert(
        "end",
        f"\nReplace complete.\n"
        f"Files with replacements: {files_with_replacements}\n"
        f"Total replacements made: {total_replacements}\n"
    )

    gui.text_results.config(state="disabled")
    gui.text_results.see("1.0")  # Scroll to top after completing
