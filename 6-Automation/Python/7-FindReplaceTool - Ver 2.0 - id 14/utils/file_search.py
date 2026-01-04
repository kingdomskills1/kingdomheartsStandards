import os
import re
from docx import Document
import fitz  # PyMuPDF for PDFs
import tkinter as tk

import os
import re
from docx import Document
import fitz  # PyMuPDF

def find_in_file(file_path, search_text, case_sensitive=False, use_regex=False, content_type="all"):
    """
    Searches for search_text in a given file and returns structured results.
    
    Returns:
        dict with:
            - 'file_path': str
            - 'matches': list of tuples (line_number, line_text, [(start, end), ...])
            - 'file_find_count': int
            - 'error': str or None
            - 'unsupported': bool
    """
    results = {
        "file_path": file_path,
        "matches": [],
        "file_find_count": 0,
        "error": None,
        "unsupported": False
    }
    
    try:
        _, ext = os.path.splitext(file_path)
        lines = []
        primary_line_count = 0
        flags = 0 if case_sensitive else re.IGNORECASE

        if ext.lower() == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                lines = [line.rstrip("\n") for line in f]
            primary_line_count = len(lines)

        elif ext.lower() == ".docx":
            doc = Document(file_path)
            primary_lines = []

            if content_type in ("all", "text"):
                for p in doc.paragraphs:
                    if p.text.strip() and (content_type != "text" or not (p.style and p.style.name.startswith("Heading"))):
                        primary_lines.append(p.text)

            if content_type in ("headings", "tables_headings"):
                for p in doc.paragraphs:
                    if p.style and p.style.name.startswith("Heading") and p.text.strip():
                        primary_lines.append(p.text)

            if content_type in ("tables", "tables_headings", "all"):
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                primary_lines.append(cell.text)

            lines = primary_lines
            primary_line_count = len(lines)

        elif ext.lower() == ".pdf":
            doc = fitz.open(file_path)
            for page in doc:
                text = page.get_text()
                if text:
                    lines.extend(text.splitlines())
            primary_line_count = len(lines)

        else:
            results["unsupported"] = True
            return results

        # --- SEARCH ---
        for i, line in enumerate(lines, start=1):
            matches = []

            if use_regex:
                try:
                    matches = [m.span() for m in re.finditer(search_text, line, flags)]
                except re.error:
                    continue
            else:
                search_line = line if case_sensitive else line.lower()
                target = search_text if case_sensitive else search_text.lower()
                pos = 0
                while True:
                    idx = search_line.find(target, pos)
                    if idx == -1:
                        break
                    matches.append((idx, idx + len(search_text)))
                    pos = idx + len(search_text)

            if matches and i <= primary_line_count:
                results["matches"].append((i, line, matches))
                results["file_find_count"] += len(matches)

    except Exception as e:
        results["error"] = str(e)

    return results


def search_path(path, search_text, selected_type="file", case_sensitive=False, use_regex=False,
                search_subfolders=False, txt_only=False, doc_only=False, pdf_only=False, content_type="all"):
    """
    Searches for text in a file or folder and returns all results as a list.
    """
    results_list = []

    extensions = []
    if txt_only: extensions.append(".txt")
    if doc_only: extensions.append(".docx")
    if pdf_only: extensions.append(".pdf")

    def file_matches(name):
        return any(name.lower().endswith(ext) for ext in extensions) if extensions else True

    if selected_type == "file" and os.path.isfile(path):
        results_list.append(find_in_file(path, search_text, case_sensitive, use_regex, content_type))

    elif selected_type == "folder" and os.path.isdir(path):
        if search_subfolders:
            for root, _, files in os.walk(path):
                for f in files:
                    if file_matches(f):
                        results_list.append(find_in_file(os.path.join(root, f), search_text, case_sensitive, use_regex, content_type))
        else:
            for f in os.listdir(path):
                full = os.path.join(path, f)
                if os.path.isfile(full) and file_matches(f):
                    results_list.append(find_in_file(full, search_text, case_sensitive, use_regex, content_type))

    return results_list


def display_results(gui, results_list):
    gui.text_results.config(state="normal")
    gui.text_results.tag_remove("highlight", "1.0", "end")
    gui.text_results.tag_remove("bold", "1.0", "end")
    gui.text_results.tag_configure("bold", font=("TkDefaultFont", 10, "bold"))
    gui.text_results.tag_configure("highlight", background="yellow")
    gui.text_results.delete("1.0", "end")

    total_finds = 0
    files_with_finds = 0

    for result in results_list:
        file_path = result["file_path"]

        if result["unsupported"]:
            start = gui.text_results.index("end-1c")
            gui.text_results.insert("end", f"Unsupported file type: {file_path}\n\n")
            end = gui.text_results.index("end-1c")
            gui.text_results.tag_add("bold", start, end)
            continue

        if result["error"]:
            start = gui.text_results.index("end-1c")
            gui.text_results.insert("end", f"Error reading file: {file_path} -> {result['error']}\n\n")
            end = gui.text_results.index("end-1c")
            gui.text_results.tag_add("bold", start, end)
            continue

        file_find_count = result["file_find_count"]
        matches = result["matches"]

        if file_find_count > 0:
            files_with_finds += 1
            total_finds += file_find_count

            start = gui.text_results.index("end-1c")
            gui.text_results.insert("end", f"{file_path} ({file_find_count} matches)\n")
            end = gui.text_results.index("end-1c")
            gui.text_results.tag_add("bold", start, end)

            for line_no, text, line_matches in matches:
                ln_start = gui.text_results.index("end-1c")
                gui.text_results.insert("end", f"line [{line_no}]: ")
                ln_end = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", ln_start, ln_end)

                content_start = gui.text_results.index("end-1c")
                gui.text_results.insert("end", text + "\n")

                for s, e in line_matches:
                    gui.text_results.tag_add(
                        "highlight",
                        f"{content_start}+{s}c",
                        f"{content_start}+{e}c"
                    )
            gui.text_results.insert("end", "\n")
        else:
            start = gui.text_results.index("end-1c")
            gui.text_results.insert("end", f"{file_path}\n")
            end = gui.text_results.index("end-1c")
            gui.text_results.tag_add("bold", start, end)
            gui.text_results.insert("end", "No matches found in this file.\n\n")

    gui.text_results.insert("end", f"Search complete.\nFiles with matches: {files_with_finds}\nTotal matches found: {total_finds}\n\n")
    gui.text_results.config(state="disabled")
    gui.text_results.see("1.0")


def run_search_file_paths_only(gui, path, search_text, selected_type="file",
                               case_sensitive=False, use_regex=False, search_subfolders=False,
                               txt_only=False, doc_only=False, pdf_only=False,
                               content_type="all"):  # <-- added content_type
    """
    Search a file or folder for text (supports regex and case sensitivity),
    and insert only the file paths of files that contain matches into gui.text_results.
    Skips unsupported files silently.
    """

    if not search_text.strip():
        from tkinter import messagebox
        messagebox.showwarning("Empty Find Field", "The Find field is empty. Please enter text to search.")
        return

    gui.text_results.config(state="normal")
    gui.text_results.delete("1.0", "end")  # Clear previous results
    gui.text_results.tag_remove("bold", "1.0", "end")
    gui.text_results.tag_configure("bold", font=("TkDefaultFont", 10, "bold"))

    def process_file(file_path):
        try:
            _, ext = os.path.splitext(file_path)
            lines = []

            if ext.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    lines = [line.rstrip("\n") for line in f if line.strip()]

            elif ext.lower() == ".docx":
                doc = Document(file_path)
                lines = []

                # Apply content_type filter
                if content_type == "all":
                    lines = [p.text for p in doc.paragraphs if p.text.strip()]
                elif content_type == "tables":
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    lines.append(cell.text)
                elif content_type == "headings":
                    for p in doc.paragraphs:
                        if p.style.name.startswith("Heading") and p.text.strip():
                            lines.append(p.text)
                elif content_type == "tables_headings":
                    for p in doc.paragraphs:
                        if p.style.name.startswith("Heading") and p.text.strip():
                            lines.append(p.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    lines.append(cell.text)

            elif ext.lower() == ".pdf":
                doc = fitz.open(file_path)
                lines = [text_line for page in doc for text_line in page.get_text().splitlines() if text_line.strip()]
            else:
                return  # skip unsupported files

            found = False
            flags = 0 if case_sensitive else re.IGNORECASE

            for line in lines:
                if use_regex:
                    try:
                        if re.search(search_text, line, flags):
                            found = True
                            break
                    except re.error:
                        continue
                else:
                    line_to_search = line if case_sensitive else line.lower()
                    search_for = search_text if case_sensitive else search_text.lower()
                    if search_for in line_to_search:
                        found = True
                        break

            if found:
                start_index = gui.text_results.index("end-1c")
                gui.text_results.insert("end", f"{file_path}\n")
                end_index = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", start_index, end_index)

        except Exception:
            pass  # silently ignore read errors

    # --- Process files/folders ---
    if selected_type == "file" and os.path.isfile(path):
        process_file(path)
    elif selected_type == "folder" and os.path.isdir(path):
        extensions = []
        if txt_only: extensions.append(".txt")
        if doc_only: extensions.append(".docx")
        if pdf_only: extensions.append(".pdf")

        def file_matches(filename):
            return any(filename.lower().endswith(ext) for ext in extensions) if extensions else True

        for root, dirs, files in os.walk(path):
            for f in files:
                file_path = os.path.join(root, f)
                if file_matches(f):
                    process_file(file_path)
            if not search_subfolders:
                break

    gui.text_results.config(state="disabled")
    gui.text_results.see("1.0")
