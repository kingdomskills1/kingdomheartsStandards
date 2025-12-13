import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import tkinter.font as tkFont
import unicodedata
from docx.enum.text import WD_COLOR_INDEX



# ----------------- Helpers -----------------
def extract_plain_text(tk_text_widget):
    """
    Extract plain text from the Replace Text widget.
    """
    return tk_text_widget.get("1.0", "end-1c")


def normalize_text(s):
    """
    Normalize Word text to handle special characters like:
    - non-breaking hyphens
    - non-breaking spaces
    - Unicode normalization
    """
    s = s.replace("\u2011", "-").replace("\u00A0", " ")
    s = unicodedata.normalize("NFKC", s)
    return s


def get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf):
    """
    Return a list of files to process.
    """
    files = []
    if is_file:
        return [path]

    allowed_exts = []
    if txt: allowed_exts.append(".txt")
    if doc: allowed_exts.append(".docx")
    if pdf: allowed_exts.append(".pdf")

    for root, dirs, filenames in os.walk(path):
        for name in filenames:
            if not allowed_exts or any(name.lower().endswith(ext) for ext in allowed_exts):
                files.append(os.path.join(root, name))
        if not include_subfolders:
            break

    return files


# ----------------- TXT Replacement -----------------
def replace_in_file_txt(path, find_text, replace_text, case_sensitive, regex):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    flags = 0 if case_sensitive else re.IGNORECASE

    if regex:
        new_content = re.sub(find_text, replace_text, content, flags=flags)
    else:
        pattern = re.compile(re.escape(find_text), flags=flags)
        new_content = pattern.sub(replace_text, content)

    with open(path, "w", encoding="utf-8", errors="ignore") as f:
        f.write(new_content)


# ----------------- DOCX Replacement -----------------
def replace_in_file_docx(path, tk_replace_widget, find_text,
                         case_sensitive=False, regex=False):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_COLOR_INDEX
    import tkinter.font as tkFont
    import re

    def normalize_text(text):
        return text.replace('\r', '').replace('\n', '')

    doc = Document(path)

    # ---------------- HEX → Word highlight name ----------------
    HEX_TO_WORD_HIGHLIGHT = {
        "#000000": "Black", "#0000ff": "Blue", "#00ffff": "Turquoise",
        "#00ff00": "Bright Green", "#ffc0cb": "Pink", "#ff0000": "Red",
        "#ffff00": "Yellow", "#ffffff": "White", "#00008b": "Dark Blue",
        "#008080": "Teal", "#008000": "Green", "#ee82ee": "Violet",
        "#8b0000": "Dark Red", "#9b870c": "Dark Yellow", "#808080": "Gray50",
        "#c0c0c0": "Gray25"
    }

    # ---------------- Word → DOCX highlight ----------------
    HIGHLIGHT_TK_TO_DOCX = {
        "Black": WD_COLOR_INDEX.BLACK, "Blue": WD_COLOR_INDEX.BLUE,
        "Turquoise": WD_COLOR_INDEX.TURQUOISE, "Bright Green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "Pink": WD_COLOR_INDEX.PINK, "Red": WD_COLOR_INDEX.RED,
        "Yellow": WD_COLOR_INDEX.YELLOW, "White": WD_COLOR_INDEX.WHITE,
        "Dark Blue": WD_COLOR_INDEX.DARK_BLUE, "Teal": WD_COLOR_INDEX.TEAL,
        "Green": WD_COLOR_INDEX.GREEN, "Violet": WD_COLOR_INDEX.VIOLET,
        "Dark Red": WD_COLOR_INDEX.DARK_RED, "Dark Yellow": WD_COLOR_INDEX.DARK_YELLOW,
        "Gray50": WD_COLOR_INDEX.GRAY_50, "Gray25": WD_COLOR_INDEX.GRAY_25,
        "None": None
    }

    # ---------------- Extract replacement text + formatting ----------------
    replace_text = tk_replace_widget.get("1.0", "end-1c")
    replace_runs = []

    for i, ch in enumerate(replace_text):
        idx = f"1.0 + {i} chars"
        tags = tk_replace_widget.tag_names(idx)

        font_family = "Arial"
        font_size = 12
        bold = False
        italic = False
        highlight_name = None

        for tag in tags:
            # Font
            try:
                f = tk_replace_widget.tag_cget(tag, "font")
                if f:
                    font_info = tkFont.Font(font=f)
                    font_family = font_info.actual("family")
                    font_size = font_info.actual("size")
                    bold = font_info.actual("weight") == "bold"
                    italic = font_info.actual("slant") == "italic"
            except:
                pass

            # Highlight (HEX → Word name)
            bg = tk_replace_widget.tag_cget(tag, "background")
            if bg:
                highlight_name = HEX_TO_WORD_HIGHLIGHT.get(bg.lower())

        replace_runs.append({
            "text": ch,
            "font_family": font_family,
            "font_size": font_size,
            "bold": bold,
            "italic": italic,
            "highlight": highlight_name
        })

    # ---------------- Normalize find text ----------------
    find_text_norm = normalize_text(find_text)

    # ---------------- Process paragraphs ----------------
    for para in doc.paragraphs:
        para_text_norm = normalize_text(para.text)
        matches = []

        if regex:
            flags = 0 if case_sensitive else re.IGNORECASE
            for m in re.finditer(find_text_norm, para_text_norm, flags=flags):
                matches.append((m.start(), m.end()))
        else:
            search_text = para_text_norm if case_sensitive else para_text_norm.lower()
            target = find_text_norm if case_sensitive else find_text_norm.lower()

            start = 0
            while True:
                idx = search_text.find(target, start)
                if idx == -1:
                    break
                matches.append((idx, idx + len(find_text_norm)))
                start = idx + len(find_text_norm)

        if not matches:
            continue

        # ---------------- Build new runs ----------------
        new_runs = []
        last_idx = 0
        for start, end in matches:
            # Add text BEFORE match
            if start > last_idx:
                new_runs.append({"text": para.text[last_idx:start]})
            # Add replacement runs instead of matched text
            new_runs.extend(replace_runs)
            last_idx = end
        # Add text AFTER last match
        if last_idx < len(para.text):
            new_runs.append({"text": para.text[last_idx:]})

        # ---------------- Replace paragraph ----------------
        # Remove all old runs
        for run in para.runs:
            run.clear()

        # Add new runs (skip empty and remove trailing newlines)
        for run_info in new_runs:
            run_text = run_info["text"].rstrip('\n').rstrip('\r')
            if run_text == "":
                continue
            run = para.add_run(run_text)

            if "font_family" in run_info:
                run.font.name = run_info["font_family"]
                run.font.size = Pt(run_info["font_size"])
                run.font.bold = run_info["bold"]
                run.font.italic = run_info["italic"]
                run.font.highlight_color = HIGHLIGHT_TK_TO_DOCX.get(run_info["highlight"])

    doc.save(path)


# ----------------- Main Replace Process -----------------
def run_replace_process(gui, path, is_file, find_text, tk_replace_widget,
                        case_sensitive, regex, include_subfolders, txt, doc, pdf):
    """
    Main function to replace text in file/folder with formatting preserved.
    """
    gui.text_results.config(state="normal")
    files = get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf)
    gui.text_results.insert("end", f"Processing {len(files)} file(s)...\n")

    for file_path in files:
        try:
            if file_path.lower().endswith(".txt"):
                plain_text = extract_plain_text(tk_replace_widget)
                replace_in_file_txt(file_path, find_text, plain_text, case_sensitive, regex)

            elif file_path.lower().endswith(".docx"):
                replace_in_file_docx(file_path, tk_replace_widget, find_text, case_sensitive, regex)

            elif file_path.lower().endswith(".pdf"):
                gui.text_results.insert("end", f"PDF replace not supported: {file_path}\n")
                continue

            gui.text_results.insert("end", f"Replaced in: {file_path}\n")

        except Exception as e:
            gui.text_results.insert("end", f"Error in {file_path}: {e}\n")

    gui.text_results.config(state="disabled")
