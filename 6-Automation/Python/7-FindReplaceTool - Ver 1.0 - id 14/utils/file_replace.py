import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import tkinter.font as tkFont
import unicodedata
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from utils.inspect_word_selection import inspect_word_selection



# ----------------- Helpers -----------------
def extract_plain_text(tk_text_widget):
    """
    Extract plain text from the Replace Text widget.
    """
    return tk_text_widget.get("1.0", "end-1c")


def normalize_text(s):
    """
    Normalize Word text to handle special characters like:
    - non-breaking hyphens
    - non-breaking spaces
    - Unicode normalization
    """
    s = s.replace("\u2011", "-").replace("\u00A0", " ")
    s = unicodedata.normalize("NFKC", s)
    return s


def get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf):
    """
    Return a list of files to process.
    """
    files = []
    if is_file:
        return [path]

    allowed_exts = []
    if txt: allowed_exts.append(".txt")
    if doc: allowed_exts.append(".docx")
    if pdf: allowed_exts.append(".pdf")

    for root, dirs, filenames in os.walk(path):
        for name in filenames:
            if not allowed_exts or any(name.lower().endswith(ext) for ext in allowed_exts):
                files.append(os.path.join(root, name))
        if not include_subfolders:
            break

    return files


# ----------------- TXT Replacement -----------------
def replace_in_file_txt(path, find_text, replace_text, case_sensitive, regex):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    flags = 0 if case_sensitive else re.IGNORECASE

    if regex:
        new_content = re.sub(find_text, replace_text, content, flags=flags)
    else:
        pattern = re.compile(re.escape(find_text), flags=flags)
        new_content = pattern.sub(replace_text, content)

    with open(path, "w", encoding="utf-8", errors="ignore") as f:
        f.write(new_content)



# start functions of replaced
# ----------------- Replacement Format -----------------
def get_replacement_format(tk_replace_widget):
    """
    Returns a dictionary containing text and its formatting from the current Word selection.
    This function directly extracts font properties (family, size, bold, italic, highlight)
    from Word selection using win32com, and ensures proper formatting during replacement in DOCX.
    """
    import win32com.client
    from docx.enum.text import WD_COLOR_INDEX

    # Mapping of highlight index to highlight names
    HIGHLIGHT_MAP = {
        0: "None", 1: "Black", 2: "Blue", 3: "Turquoise",
        4: "Bright Green", 5: "Pink", 6: "Red", 7: "Yellow", 8: "White",
        9: "Dark Blue", 10: "Teal", 11: "Green", 12: "Violet",
        13: "Dark Red", 14: "Dark Yellow", 15: "Gray50", 16: "Gray25"
    }

    # Mapping highlight names to actual color codes (supported by python-docx)
    WORD_TO_COLOR = {
        "Black": WD_COLOR_INDEX.BLACK, "Blue": WD_COLOR_INDEX.BLUE,
        "Turquoise": WD_COLOR_INDEX.TURQUOISE, "Bright Green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "Pink": WD_COLOR_INDEX.PINK, "Red": WD_COLOR_INDEX.RED, "Yellow": WD_COLOR_INDEX.YELLOW,
        "White": WD_COLOR_INDEX.WHITE, "Dark Blue": WD_COLOR_INDEX.DARK_BLUE, 
        "Teal": WD_COLOR_INDEX.TEAL, "Green": WD_COLOR_INDEX.GREEN, 
        "Violet": WD_COLOR_INDEX.VIOLET, "Dark Red": WD_COLOR_INDEX.DARK_RED,
        "Dark Yellow": WD_COLOR_INDEX.DARK_YELLOW, 
        "Gray50": WD_COLOR_INDEX.GRAY_50, 
        "Gray25": WD_COLOR_INDEX.GRAY_25,  # Map Gray25 to Gray (could be changed)
        "None": None  # None value for no highlight
    }

    # If the widget is empty, don't proceed
    text_content = tk_replace_widget.get("1.0", "end-1c").strip()  # Get all content in Text widget
    if not text_content:
        return None  # Return None if there's no selected text to replace

    # Proceed with getting the font formatting as before
    word = win32com.client.Dispatch("Word.Application")
    rng = word.Selection.Range
    first_char = rng.Characters(1)
    font = first_char.Font

    font_name = font.Name or "Arial"
    font_size = int(font.Size) if font.Size else 12
    bold = bool(font.Bold)
    italic = bool(font.Italic)

    # Get the highlight color
    try:
        idx = first_char.HighlightColorIndex
        highlight_name = HIGHLIGHT_MAP.get(idx, "None")
        print(f"Raw HighlightColorIndex: {idx}, Highlight Name: {highlight_name}")
    except Exception as e:
        print(f"Error getting highlight color: {e}")
        highlight_name = "None"

    # Get the corresponding WD_COLOR_INDEX highlight color
    highlight_color = WORD_TO_COLOR.get(highlight_name, WD_COLOR_INDEX.YELLOW)

    return {
        "text": text_content,
        "font_family": font_name,
        "font_size": font_size,
        "bold": bold,
        "italic": italic,
        "highlight": highlight_color
    }


# ----------------- Safe Paragraph Replacement -----------------
def replace_paragraph_safe(para, pattern, repl):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if not para.text:
        return

    original_text = para.text
    matches = list(pattern.finditer(original_text))
    if not matches:
        return

    # Save paragraph style and font from the first run
    para_style = para.style
    default_font = None
    default_size = None
    default_bold = None
    default_italic = None
    default_highlight = None

    if para.runs:
        r = para.runs[0]
        default_font = r.font.name
        default_size = r.font.size
        default_bold = r.font.bold
        default_italic = r.font.italic
        default_highlight = r.font.highlight_color

    # Clear the paragraph content but keep its style
    para.clear()
    para.style = para_style

    # Function to force font family in DOCX XML
    def force_run_font(run, font_name):
        """Force the run to use a specific font in DOCX (works in Word)."""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
        run.font.name = font_name  # also set in python-docx object

    # Track position in original text
    pos = 0
    for m in matches:
        start, end = m.span()

        # Text before the match (unchanged text)
        if start > pos:
            r = para.add_run(original_text[pos:start])
            r.font.name = default_font
            r.font.size = default_size
            r.font.bold = default_bold
            r.font.italic = default_italic
            if default_highlight:
                r.font.highlight_color = default_highlight

        # Replacement text (highlight and font changes)
        r = para.add_run(repl["text"])
        r.font.size = Pt(repl["font_size"])
        r.font.bold = repl["bold"]
        r.font.italic = repl["italic"]
        if repl["highlight"]:
            r.font.highlight_color = repl["highlight"]
        force_run_font(r, repl["font_family"])

        pos = end

    # Remaining text after the last match (unchanged text)
    if pos < len(original_text):
        r = para.add_run(original_text[pos:])
        r.font.name = default_font
        r.font.size = default_size
        r.font.bold = default_bold
        r.font.italic = default_italic
        if default_highlight:
            r.font.highlight_color = default_highlight


# ----------------- Replacement When Empty -----------------
def replace_when_empty_safe(para, pattern):
    if not para.text:
        return

    new_text = pattern.sub("", para.text)
    if new_text == para.text:
        return

    style = para.style
    para.clear()
    para.style = style
    para.add_run(new_text)

# ----------------- Tables Support -----------------
def process_tables(tables, handler):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    handler(para)
                if cell.tables:
                    process_tables(cell.tables, handler)

# ----------------- DOCX Replacement -----------------
def replace_in_file_docx(path, tk_replace_widget, find_text,
                         case_sensitive=False, regex=False):
    from docx import Document
    import re

    # Load the document
    doc = Document(path)
    
    # Get replacement format
    repl = get_replacement_format(tk_replace_widget)

    # Check if replacement format is valid
    if repl is None:
        print(f"Skipping replacement for {path}: No text selected.")
        return  # No valid replacement format, exit early

    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(find_text, flags) if regex else re.compile(re.escape(find_text), flags)

    # If there's no text to replace (empty text), we still need to process with empty replacement
    handler = lambda p: replace_paragraph_safe(p, pattern, repl) if repl["text"] else lambda p: replace_when_empty_safe(p, pattern)

    # Process paragraphs in the document
    for para in doc.paragraphs:
        handler(para)

    # Process tables if any
    process_tables(doc.tables, handler)

    # Save the document after replacement
    doc.save(path)
    print(f"Replaced text in {path}.")


# end functions of replaced

# ----------------- Main Replace Process -----------------
def run_replace_process(gui, path, is_file, find_text, tk_replace_widget,
                        case_sensitive, regex, include_subfolders, txt, doc, pdf):
    
    # inspect_word_selection()
    """
    Main function to replace text in file/folder with formatting preserved.
    """
    gui.text_results.config(state="normal")
    files = get_files_to_process(path, is_file, include_subfolders, txt, doc, pdf)
    gui.text_results.insert("end", f"Processing {len(files)} file(s)...\n")

    for file_path in files:
        try:
            if file_path.lower().endswith(".txt"):
                plain_text = extract_plain_text(tk_replace_widget)
                replace_in_file_txt(file_path, find_text, plain_text, case_sensitive, regex)

            elif file_path.lower().endswith(".docx"):
                replace_in_file_docx(file_path, tk_replace_widget, find_text, case_sensitive, regex)
                # replace_in_file_docx(file_path, tk_replace_widget, find_text, case_sensitive, regex)

            elif file_path.lower().endswith(".pdf"):
                gui.text_results.insert("end", f"PDF replace not supported: {file_path}\n")
                continue

            gui.text_results.insert("end", f"Replaced in: {file_path}\n")

        except Exception as e:
            gui.text_results.insert("end", f"Error in {file_path}: {e}\n")

    gui.text_results.config(state="disabled")
