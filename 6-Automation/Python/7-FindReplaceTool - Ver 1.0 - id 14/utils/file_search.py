import os
import re
from docx import Document
import fitz  # PyMuPDF for PDFs
import tkinter as tk

def run_search_in_path(gui, path, search_text, selected_type="file",
                       case_sensitive=False, use_regex=False, search_subfolders=False,
                       txt_only=False, doc_only=False, pdf_only=False, silent_mode=False):
    """
    Search a file or folder for text (supports regex and case sensitivity),
    and insert results into the provided Tkinter Text widget (gui.text_results).
    Supports .txt, .docx, and .pdf files.

    silent_mode=True -> Search Match Only: show only files with matches.
    silent_mode=False -> Run Search: show all files, including no matches and unsupported types.
    """

    gui.text_results.config(state="normal")
    gui.text_results.tag_remove("highlight", "1.0", "end")
    gui.text_results.tag_remove("bold", "1.0", "end")
    gui.text_results.delete("1.0", "end")

    gui.text_results.tag_configure("bold", font=("TkDefaultFont", 10, "bold"))
    gui.text_results.tag_configure("highlight", background="yellow")

    def process_file(file_path):
        lines = []
        output_exists = False  # Tracks if anything will be printed for this file

        try:
            _, ext = os.path.splitext(file_path)
            if ext.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    lines = [line.rstrip("\n") for line in f if line.strip()]
            elif ext.lower() == ".docx":
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
            elif ext.lower() == ".pdf":
                doc = fitz.open(file_path)
                for page in doc:
                    text = page.get_text().strip()
                    if text:
                        lines.extend(text.splitlines())
            else:
                if not silent_mode:  # Show unsupported only in Run Search
                    start_index = gui.text_results.index("end-1c")
                    gui.text_results.insert("end", f"Unsupported file type: {ext} ({file_path})\n\n")
                    end_index = gui.text_results.index("end-1c")
                    gui.text_results.tag_add("bold", start_index, end_index)        
                    output_exists = True
                return
        except Exception as e:
            if not silent_mode:
                start_index = gui.text_results.index("end-1c")
                gui.text_results.insert("end", f"Error reading file: {file_path} -> {e}\n\n")
                end_index = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", start_index, end_index)    
                output_exists = True
            return

        found_in_file = False
        file_path_inserted = False
        flags = 0 if case_sensitive else re.IGNORECASE

        for i, line in enumerate(lines, start=1):
            matches = []

            if use_regex:
                try:
                    matches = [m.span() for m in re.finditer(search_text, line, flags)]
                except re.error:
                    continue
            else:
                line_to_search = line if case_sensitive else line.lower()
                search_for = search_text if case_sensitive else search_text.lower()
                start_pos = 0
                while True:
                    idx = line_to_search.find(search_for, start_pos)
                    if idx == -1:
                        break
                    matches.append((idx, idx + len(search_text)))
                    start_pos = idx + len(search_text)

            if matches:
                if not file_path_inserted:
                    start_index = gui.text_results.index("end-1c")
                    gui.text_results.insert("end", f"{file_path}\n")
                    end_index = gui.text_results.index("end-1c")
                    gui.text_results.tag_add("bold", start_index, end_index)
                    file_path_inserted = True

                found_in_file = True
                output_exists = True

                # Insert line number
                line_number_text = f"line [{i}]: "
                start_ln = gui.text_results.index("end-1c")
                gui.text_results.insert("end", line_number_text)
                end_ln = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", start_ln, end_ln)

                # Insert content
                content_start_index = gui.text_results.index("end-1c")
                gui.text_results.insert("end", line + "\n")
                content_end_index = gui.text_results.index("end-1c")

                # Highlight matches
                for start_idx, end_idx in matches:
                    match_start = f"{content_start_index}+{start_idx}c"
                    match_end = f"{content_start_index}+{end_idx}c"
                    gui.text_results.tag_add("highlight", match_start, match_end)

        # Handle files with no matches
        if not found_in_file and not silent_mode:
            if not file_path_inserted:
                start_index = gui.text_results.index("end-1c")
                gui.text_results.insert("end", f"{file_path}\n")
                end_index = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", start_index, end_index)   
                file_path_inserted = True
            gui.text_results.insert("end", "No matches found in this file.\n")
            output_exists = True

        # Add one blank line only if this file produced output
        if output_exists:
            gui.text_results.insert("end", "\n")

    # --- Process file or folder ---
    if selected_type == "file":
        if os.path.isfile(path):
            process_file(path)
    elif selected_type == "folder":
        if os.path.isdir(path):
            extensions = []
            if txt_only:
                extensions.append(".txt")
            if doc_only:
                extensions.append(".docx")
            if pdf_only:
                extensions.append(".pdf")

            def file_matches(filename):
                return any(filename.lower().endswith(ext) for ext in extensions) if extensions else True

            if search_subfolders:
                for root, dirs, files in os.walk(path):
                    for f in files:
                        if file_matches(f):
                            process_file(os.path.join(root, f))
            else:
                for f in os.listdir(path):
                    file_path = os.path.join(path, f)
                    if os.path.isfile(file_path) and file_matches(f):
                        process_file(file_path)

    gui.text_results.config(state="disabled")
    gui.text_results.see("1.0")



def run_search_file_paths_only(gui, path, search_text, selected_type="file",
                               case_sensitive=False, use_regex=False, search_subfolders=False,
                               txt_only=False, doc_only=False, pdf_only=False):
    """
    Search a file or folder for text (supports regex and case sensitivity),
    and insert only the file paths of files that contain matches into gui.text_results.
    Skips unsupported files silently.
    """

    if not search_text.strip():
        from tkinter import messagebox
        messagebox.showwarning("Empty Find Field", "The Find field is empty. Please enter text to search.")
        return

    gui.text_results.config(state="normal")
    gui.text_results.delete("1.0", "end")  # Clear previous results
    gui.text_results.tag_remove("bold", "1.0", "end")

    gui.text_results.tag_configure("bold", font=("TkDefaultFont", 10, "bold"))

    def process_file(file_path):
        try:
            _, ext = os.path.splitext(file_path)
            lines = []
            if ext.lower() == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    lines = [line.rstrip("\n") for line in f if line.strip()]
            elif ext.lower() == ".docx":
                doc = Document(file_path)
                lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            elif ext.lower() == ".pdf":
                doc = fitz.open(file_path)
                lines = [text_line for page in doc for text_line in page.get_text().splitlines() if text_line.strip()]
            else:
                return  # skip unsupported files

            found = False
            flags = 0 if case_sensitive else re.IGNORECASE

            for line in lines:
                if use_regex:
                    try:
                        if re.search(search_text, line, flags):
                            found = True
                            break
                    except re.error:
                        continue
                else:
                    line_to_search = line if case_sensitive else line.lower()
                    search_for = search_text if case_sensitive else search_text.lower()
                    if search_for in line_to_search:
                        found = True
                        break

            if found:
                start_index = gui.text_results.index("end-1c")
                gui.text_results.insert("end", f"{file_path}\n")
                end_index = gui.text_results.index("end-1c")
                gui.text_results.tag_add("bold", start_index, end_index)

        except Exception:
            pass  # silently ignore read errors

    # --- Process files/folders ---
    if selected_type == "file" and os.path.isfile(path):
        process_file(path)
    elif selected_type == "folder" and os.path.isdir(path):
        extensions = []
        if txt_only: extensions.append(".txt")
        if doc_only: extensions.append(".docx")
        if pdf_only: extensions.append(".pdf")

        def file_matches(filename):
            return any(filename.lower().endswith(ext) for ext in extensions) if extensions else True

        for root, dirs, files in os.walk(path):
            for f in files:
                file_path = os.path.join(root, f)
                if file_matches(f):
                    process_file(file_path)
            if not search_subfolders:
                break

    gui.text_results.config(state="disabled")
    gui.text_results.see("1.0")
