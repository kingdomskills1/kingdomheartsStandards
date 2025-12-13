import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import pyodbc
import pandas as pd
from difflib import get_close_matches
from docx import Document

class TVD_GUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Table Value Dictionary Creator")
        self.conn = None
        self.db_path = None
        self.table_name = None
        self.selected_columns = {}
        self.label_column_map = {}
        self.lamp_labels = {}

        # Database selection
        tk.Button(root, text="Select Access File", command=self.select_db).pack(pady=5)
        tk.Button(root, text="Reset Selections", command=self.reset).pack(pady=5)

        # Table selection
        tk.Label(root, text="Tables in Database:").pack()
        self.table_listbox = tk.Listbox(root, width=50, height=6)
        self.table_listbox.pack()
        self.table_listbox.bind("<Double-Button-1>", self.copy_table_name)
        self.table_listbox.bind("<<ListboxSelect>>", self.show_columns)

        # Scrollable Columns Frame with border
        tk.Label(root, text="Columns (Select & Include row_range if needed):").pack(pady=(10,0))
        self.column_container = tk.Frame(root, bd=2, relief="groove")
        self.column_container.pack(fill="both", expand=True, padx=5, pady=5)

        self.column_canvas = tk.Canvas(self.column_container, height=250)
        self.column_canvas.pack(side="left", fill="both", expand=True)

        self.column_scrollbar = tk.Scrollbar(self.column_container, orient="vertical", command=self.column_canvas.yview)
        self.column_scrollbar.pack(side="right", fill="y")
        self.column_canvas.configure(yscrollcommand=self.column_scrollbar.set)

        self.column_frame = tk.Frame(self.column_canvas)
        self.column_canvas.create_window((0, 0), window=self.column_frame, anchor='nw')
        self.column_frame.bind("<Configure>", lambda e: self.column_canvas.configure(scrollregion=self.column_canvas.bbox("all")))

        # Select All row_range checkbox
        self.select_all_var = tk.BooleanVar()
        self.select_all_chk = tk.Checkbutton(root, text="Select All row_range", variable=self.select_all_var,
                                             command=self.toggle_all_row_range)
        self.select_all_chk.pack(pady=5)

        # Create TVD file button
        self.create_btn = tk.Button(root, text="Create TVD File", command=self.create_file)
        self.create_btn.pack(pady=10)

    def select_db(self):
        self.db_path = filedialog.askopenfilename(title="Select Access Database",
                                                  filetypes=[("Access Files", "*.accdb *.mdb")])
        if not self.db_path:
            return
        try:
            conn_str = (
                r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};'
                fr'DBQ={self.db_path};'
            )
            self.conn = pyodbc.connect(conn_str)
            cursor = self.conn.cursor()
            self.tables = [row.table_name for row in cursor.tables(tableType='TABLE')]
            self.table_listbox.delete(0, tk.END)
            for t in self.tables:
                self.table_listbox.insert(tk.END, t)
            messagebox.showinfo("Success", "Database connected!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to connect to database:\n{e}")

    def copy_table_name(self, event):
        selection = self.table_listbox.curselection()
        if selection:
            table_name = self.table_listbox.get(selection[0])
            self.root.clipboard_clear()
            self.root.clipboard_append(table_name)
            messagebox.showinfo("Copied", f"Table name '{table_name}' copied to clipboard!")

    def show_columns(self, event):
        selection = self.table_listbox.curselection()
        if not selection:
            return
        self.table_name = self.table_listbox.get(selection[0])
        cursor = self.conn.cursor()
        cursor.execute(f"SELECT * FROM [{self.table_name}] WHERE 1=0")
        columns = [column[0] for column in cursor.description]

        for widget in self.column_frame.winfo_children():
            widget.destroy()
        self.selected_columns.clear()
        self.label_column_map.clear()
        self.lamp_labels.clear()
        self.select_all_var.set(False)

        for col in columns:
            col_var = tk.BooleanVar()
            row_range_var = tk.BooleanVar()
            frame = tk.Frame(self.column_frame)
            frame.pack(fill="x", anchor="w", pady=1)

            col_chk = tk.Checkbutton(frame, text=col, variable=col_var,
                                     command=lambda c=col: self.suggest_label_column(c))
            col_chk.pack(side="left")

            row_chk = tk.Checkbutton(frame, text="Include row_range", variable=row_range_var)
            row_chk.pack(side="left", padx=5)

            copy_btn = tk.Button(frame, text="Copy Name", command=lambda c=col: self.copy_column_name(c), width=10)
            copy_btn.pack(side="left", padx=5)

            lamp = tk.Label(frame, text="●", fg="grey", font=("Arial", 12))
            lamp.pack(side="left", padx=5)
            self.lamp_labels[col] = lamp

            self.selected_columns[col] = {"selected": col_var, "row_range": row_range_var, "lamp": lamp}

            # Update lamp color based on uniqueness
            self.update_lamp(col)

    def copy_column_name(self, column):
        self.root.clipboard_clear()
        self.root.clipboard_append(column)
        messagebox.showinfo("Copied", f"Column name '{column}' copied to clipboard!")

    def suggest_label_column(self, column):
        if not self.selected_columns[column]["selected"].get():
            self.label_column_map[column] = None
            return

        cursor = self.conn.cursor()
        cursor.execute(f"SELECT * FROM [{self.table_name}] WHERE 1=0")
        text_columns = [c[0] for c in cursor.description if c[1] in (12, 202, 203)]
        if column in text_columns:
            text_columns.remove(column)

        matches = get_close_matches(column, text_columns, n=3, cutoff=0.3)
        prompt = "Suggested label columns:\n" + ", ".join(matches) + "\n\nEnter label column (optional):"
        label_col = simpledialog.askstring("Label Column", prompt)
        self.label_column_map[column] = label_col if label_col else None

        # After assigning label, update lamp
        self.update_lamp(column)

    def update_lamp(self, column):
        """Check if each value has unique label and update lamp color"""
        lamp = self.selected_columns[column]["lamp"]
        label_col = self.label_column_map.get(column)
        if not label_col:
            lamp.config(fg="grey")
            return

        try:
            query = f"SELECT [{column}], [{label_col}] FROM [{self.table_name}]"
            df = pd.read_sql(query, self.conn)
            grouped = df.groupby(column)[label_col].nunique()
            if (grouped > 1).any():
                lamp.config(fg="red")  # more than one label per value
            else:
                lamp.config(fg="green")  # unique labels
        except Exception as e:
            lamp.config(fg="grey")

    def toggle_all_row_range(self):
        state = self.select_all_var.get()
        for info in self.selected_columns.values():
            info["row_range"].set(state)

    def create_file(self):
        if not self.table_name:
            messagebox.showerror("Error", "Please select a table.")
            return
        selected_cols = [col for col, info in self.selected_columns.items() if info["selected"].get()]
        if not selected_cols:
            messagebox.showerror("Error", "Please select at least one column.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")],
            title="Save TVD File As"
        )
        if not file_path:
            return

        if file_path.endswith(".txt"):
            self.save_txt(file_path, selected_cols)
        else:
            self.save_doc(file_path, selected_cols)

        messagebox.showinfo("Done", f"TVD file created at:\n{file_path}")

    def save_txt(self, file_path, selected_cols):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"# Table: {self.table_name}\n\n")
            for col in selected_cols:
                include_row_range = self.selected_columns[col]["row_range"].get()
                label_column = self.label_column_map.get(col)

                f.write(f"## Column: {col}\n")
                f.write("- Type: string\n")
                f.write("- Allowed values:\n")

                query = f"SELECT DISTINCT [{col}]" + (f", [{label_column}]" if label_column else "") + f" FROM [{self.table_name}]"
                df = pd.read_sql(query, self.conn)

                for idx, row in df.iterrows():
                    f.write(f"  - Value: {row[col]}\n")
                    f.write(f"    Label: {row[label_column] if label_column else ''}\n")
                    f.write(f"    Description: \n")

                    if include_row_range:
                        rows_df = pd.read_sql(f"SELECT * FROM [{self.table_name}]", self.conn)
                        row_indices = rows_df.index[rows_df[col] == row[col]].tolist()
                        f.write(f"    Row_range: {self.format_row_range(row_indices)}\n")
                    else:
                        f.write(f"    Row_range: \n")

                f.write("\n")

    def save_doc(self, file_path, selected_cols):
        doc = Document()

        # Heading for table
        doc.add_heading(f"Table: {self.table_name}", level=1)

        for col in selected_cols:
            include_row_range = self.selected_columns[col]["row_range"].get()
            label_column = self.label_column_map.get(col)

            # Heading for column
            doc.add_heading(f"Column: {col}", level=2)
            doc.add_paragraph("Type: string")
            doc.add_paragraph("Allowed values:")

            query = f"SELECT DISTINCT [{col}]" + (f", [{label_column}]" if label_column else "") + f" FROM [{self.table_name}]"
            df = pd.read_sql(query, self.conn)

            for idx, row in df.iterrows():
                p_val = doc.add_paragraph(style='List Bullet')
                p_val.add_run("Value: ").bold = True
                p_val.add_run(f"{row[col]}")

                p_label = doc.add_paragraph(style='List Bullet 2')
                p_label.add_run("Label: ").bold = True
                p_label.add_run(f"{row[label_column] if label_column else ''}")

                p_desc = doc.add_paragraph(style='List Bullet 2')
                p_desc.add_run("Description: ").bold = True
                p_desc.add_run("")

                p_row = doc.add_paragraph(style='List Bullet 2')
                p_row.add_run("Row_range: ").bold = True

                if include_row_range:
                    rows_df = pd.read_sql(f"SELECT * FROM [{self.table_name}]", self.conn)
                    row_indices = rows_df.index[rows_df[col] == row[col]].tolist()
                    p_row.add_run(f"{self.format_row_range(row_indices)}")
                else:
                    p_row.add_run("")

        doc.save(file_path)

    @staticmethod
    def format_row_range(indices):
        if not indices:
            return ""
        indices = sorted([i+1 for i in indices])
        ranges = []
        start = prev = indices[0]
        for i in indices[1:]:
            if i == prev + 1:
                prev = i
            else:
                ranges.append(f"{start}:{prev}" if start != prev else f"{start}")
                start = prev = i
        ranges.append(f"{start}:{prev}" if start != prev else f"{start}")
        return ",".join(ranges)

    def reset(self):
        self.table_listbox.selection_clear(0, tk.END)
        for widget in self.column_frame.winfo_children():
            widget.destroy()
        self.selected_columns.clear()
        self.label_column_map.clear()
        self.lamp_labels.clear()
        self.table_name = None
        self.select_all_var.set(False)
        messagebox.showinfo("Reset", "Selections have been cleared.")

if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("550x600")
    app = TVD_GUI(root)
    root.mainloop()
