import os
import re
import pyodbc
from docx import Document
from docx.shared import Pt
from tkinter import Tk, filedialog, simpledialog

# --- Folder selection ---
def choose_folder():
    Tk().withdraw()
    folder_path = filedialog.askdirectory(title="Select Main Folder Containing Database Folders")
    return folder_path

# --- Connect to Access DB with optional password ---
def connect_access_db(db_path):
    while True:
        try:
            conn_str = (
                r"DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};"
                f"DBQ={db_path};"
            )
            conn = pyodbc.connect(conn_str, autocommit=True)
            return conn
        except pyodbc.Error as e:
            # Check if password required
            error_msg = str(e)
            if "Not a valid password" in error_msg:
                # Ask for password via Tkinter popup
                Tk().withdraw()
                password = simpledialog.askstring(
                    "Database Password",
                    f"Enter password for database:\n{db_path}\nLeave blank to skip:",
                    show="*"
                )
                if not password:
                    print(f"Skipping {db_path} due to missing password.")
                    return None
                try:
                    conn_str = (
                        r"DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};"
                        f"DBQ={db_path};"
                        f"PWD={password};"
                    )
                    conn = pyodbc.connect(conn_str, autocommit=True)
                    return conn
                except Exception as e2:
                    print(f"❌ Could not connect with password: {e2}")
                    # Loop again to ask password
            else:
                print(f"❌ Could not connect to {db_path}: {e}")
                return None

# --- Get table names from DB ---
def get_table_names(cursor):
    table_names = []

    try:
        cursor.execute(
            "SELECT Name FROM MSysObjects "
            "WHERE Type = 1 AND Flags = 0 AND Name NOT LIKE 'MSys%'"
        )
        table_names = [row[0] for row in cursor.fetchall()]
    except:
        pass

    if not table_names:
        try:
            for row in cursor.tables(tableType='TABLE'):
                name = row.table_name
                if not name.startswith("MSys"):
                    table_names.append(name)
        except:
            pass

    filtered = []
    for name in table_names:
        if name.startswith("~TMP") or "TMPCLP" in name.upper():
            continue
        filtered.append(name)

    return filtered

# --- Get table info ---
def get_table_info(connection):
    cursor = connection.cursor()
    tables = []

    table_names = get_table_names(cursor)

    for table_name in table_names:
        # --- Get columns ---
        try:
            cursor.execute(f"SELECT * FROM [{table_name}] WHERE 1=0")
            columns = [col[0] for col in cursor.description]
            col_count = len(columns)
        except Exception as e:
            print(f"⚠ Could not read columns for {table_name}: {e}")
            columns = []
            col_count = 0

        # --- Count meaningful rows ---
        try:
            non_id_columns = [c for c in columns if c.lower() != "id"]

            conditions = []
            for col in non_id_columns:
                conditions.append(
                    f"(IIf(IsNull([{col}]), '', CStr([{col}])) <> '' "
                    f"AND IIf(IsNull([{col}]), '', CStr([{col}])) <> '0')"
                )

            if conditions:
                where_clause = " OR ".join(conditions)
                query = f"SELECT COUNT(*) FROM [{table_name}] WHERE {where_clause}"
            else:
                query = f"SELECT COUNT(*) FROM [{table_name}]"

            row_count = cursor.execute(query).fetchone()[0]

        except Exception as e:
            print(f"⚠ Row count issue in {table_name}: {e}")
            row_count = 0

        tables.append((table_name, row_count, col_count))

    return tables


# --- Shorten long names ---
def shorten_name(name, max_length):
    if len(name) <= max_length:
        return name
    part_len = max_length // 2 - 2
    return name[:part_len] + "..." + name[-part_len:]

# --- Generate Word document ---
def generate_doc(output_path, data_entries):
    doc = Document()

    # Main heading
    doc.add_heading("Data Collections", level=1)

    # Calculate totals
    total_tables = total_rows = total_cols = 0
    for folder_tables in data_entries.values():
        for _, _, rows, cols in folder_tables:
            total_tables += 1
            total_rows += rows
            total_cols += cols

    # Summary
    summary_lines = [
        f"Total Tables: {total_tables}",
        f"Total Rows: {total_rows}",
        f"Total Columns: {total_cols}"
    ]
    for line in summary_lines:
        p = doc.add_paragraph(line)
        p.bold = True

    doc.add_paragraph("\n")

    # Folder-wise details
    for folder_name, folder_tables in data_entries.items():
        doc.add_heading(f"{folder_name}", level=2)

        for db_name, table_name, rows, cols in folder_tables:
            full_text = (
                f"Data Path: {folder_name} Folder -> {db_name} Database -> "
                f"{table_name} Table [Rows x Columns: {rows} x {cols}]"
            )

            # Adjust font size
            if len(full_text) >= 116:
                font_size = 7
            elif len(full_text) >= 104:
                font_size = 8    
            elif len(full_text) >= 96:
                font_size = 9
            else:
                font_size = 10

            p = doc.add_paragraph()
            parts = [
                ("Data Path: ", True),
                (f"{folder_name} Folder -> ", False),
                (f"{db_name} Database -> ", False),
                (f"{table_name} Table ", True),
                (f"[Rows x Columns: ", True),
                (f"{rows} x {cols}", False),
                ("]", True)
            ]
            for text, bold in parts:
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(font_size)

        doc.add_paragraph("\n")

    doc.save(output_path)

# --- Numeric sort ---
def sort_folders_numeric(folders):
    def extract_number(name):
        match = re.match(r'(\d+)', name)
        return int(match.group(1)) if match else float('inf')
    return sorted(folders, key=extract_number)

# --- Main function ---
def main():
    print("Choose the main folder containing all database folders...")
    main_folder = choose_folder()
    if not main_folder:
        print("No folder selected. Exiting.")
        return

    data_entries = {}

    folders = sort_folders_numeric([f for f in os.listdir(main_folder) if os.path.isdir(os.path.join(main_folder, f))])

    for folder in folders:
        folder_path = os.path.join(main_folder, folder)
        all_tables_in_folder = []

        for file in os.listdir(folder_path):
            if file.endswith(".accdb") or file.endswith(".mdb"):
                db_path = os.path.join(folder_path, file)
                db_name = os.path.splitext(file)[0]

                print(f"Reading: {db_path}")
                conn = connect_access_db(db_path)
                if conn:
                    try:
                        table_info = get_table_info(conn)
                    except Exception as e:
                        print(f"❌ Error reading tables in {db_name}: {e}")
                        table_info = []
                    conn.close()

                    for table_name, rows, cols in table_info:
                        all_tables_in_folder.append((db_name, table_name, rows, cols))

        if all_tables_in_folder:
            data_entries[folder] = all_tables_in_folder

    output_file = os.path.join(main_folder, "Data Collections.docx")
    generate_doc(output_file, data_entries)

    print("\n✅ Completed! Document saved as:")
    print(output_file)

if __name__ == "__main__":
    main()
