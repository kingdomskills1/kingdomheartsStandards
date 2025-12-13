"""Core functions for extracting and numbering headings from .docx files.

This module is separate so GUI and CLI can both import it without circular imports.
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Tuple
import argparse

from docx import Document


def extract_headings_from_docx(path: Path) -> List[Tuple[int, str]]:
    doc = Document(path)
    headings: List[Tuple[int, str]] = []
    for p in doc.paragraphs:
        style = p.style
        if style is None:
            continue
        name = getattr(style, "name", "")
        if not name:
            continue
        if name.startswith("Heading"):
            parts = name.split()
            level = 1
            if len(parts) >= 2 and parts[-1].isdigit():
                try:
                    level = int(parts[-1])
                except Exception:
                    level = 1
            headings.append((level, p.text.strip()))
    return headings


def number_headings(headings: List[Tuple[int, str]]) -> List[Tuple[str, str]]:
    if not headings:
        return []
    max_level = max(level for level, _ in headings)
    counters = [0] * max_level
    numbered: List[Tuple[str, str]] = []
    for level, text in headings:
        if level < 1:
            level = 1
        if level > len(counters):
            counters.extend([0] * (level - len(counters)))
        counters[level - 1] += 1
        for i in range(level, len(counters)):
            counters[i] = 0
        number_parts = [str(counters[i]) for i in range(level) if counters[i] != 0]
        number_str = ".".join(number_parts)
        numbered.append((number_str, text))
    return numbered


def write_headings_text(out_path: Path, numbered: List[Tuple[str, str]]) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8") as f:
        for num, text in numbered:
            f.write(f"{num} {text}\n")


def write_numbered_docx(original: Path, out_path: Path, numbered: List[Tuple[str, str]]) -> None:
    doc = Document(original)
    heading_iter = iter(numbered)
    next_expected = None
    try:
        next_expected = next(heading_iter)
    except StopIteration:
        next_expected = None

    new_doc = Document()
    for p in doc.paragraphs:
        style = p.style
        name = getattr(style, "name", "") if style is not None else ""
        text = p.text
        if name.startswith("Heading") and next_expected is not None:
            num, heading_text = next_expected
            prefixed = f"{num} {text}" if text else f"{num} {heading_text}"
            new_doc.add_paragraph(prefixed, style=name)
            if text == heading_text:
                try:
                    next_expected = next(heading_iter)
                except StopIteration:
                    next_expected = None
        else:
            if name:
                new_doc.add_paragraph(text, style=name)
            else:
                new_doc.add_paragraph(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    new_doc.save(out_path)


def process_file(path: Path, out_dir: Path, write_docx: bool) -> None:
    headings = extract_headings_from_docx(path)
    numbered = number_headings(headings)
    base = path.stem
    text_out = out_dir / (base + "_headings.txt")
    write_headings_text(text_out, numbered)
    if write_docx:
        docx_out = out_dir / (base + "_numbered.docx")
        write_numbered_docx(path, docx_out, numbered)


def process_path(src: Path, out_dir: Path, recursive: bool, pattern: str, write_docx: bool) -> None:
    if src.is_file():
        process_file(src, out_dir, write_docx)
        return
    if recursive:
        matches = src.rglob(pattern)
    else:
        matches = src.glob(pattern)
    for p in matches:
        if p.is_file():
            process_file(p, out_dir / p.parent.relative_to(src), write_docx)


def main(argv=None) -> None:
    parser = argparse.ArgumentParser(description="Copy and number heading names from .docx files")
    parser.add_argument("--src", required=True, help="Source .docx file or directory")
    parser.add_argument("--out", default="output", help="Output directory")
    parser.add_argument("--recursive", action="store_true", help="Search directories recursively")
    parser.add_argument("--pattern", default="*.docx", help="Glob pattern for files")
    parser.add_argument("--write-docx", action="store_true", help="Also write a numbered .docx copy")
    args = parser.parse_args(argv)

    src = Path(args.src)
    out = Path(args.out)
    if not src.exists():
        raise SystemExit(f"Source not found: {src}")
    process_path(src, out, args.recursive, args.pattern, args.write_docx)


if __name__ == "__main__":
    main()
