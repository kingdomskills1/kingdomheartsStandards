"""Core functions for extracting and numbering headings from .docx files."""
from __future__ import annotations

from pathlib import Path
from typing import List, Tuple
import argparse
import re
from docx import Document


def extract_headings_from_docx(path: Path) -> List[Tuple[int, str]]:
    doc = Document(path)
    headings: List[Tuple[int, str]] = []
    for p in doc.paragraphs:
        style = p.style
        if style is None:
            continue
        name = getattr(style, "name", "")
        if not name:
            continue
        if name.startswith("Heading"):
            parts = name.split()
            level = 1
            if len(parts) >= 2 and parts[-1].isdigit():
                level = int(parts[-1])
            headings.append((level, p.text.strip()))
    return headings


def clean_heading_text(text: str) -> str:
    text = re.sub(r'^\s*\d+(?:\.\d+)*\s*-\s*', '', text)
    text = re.sub(r'[^\w\s\/\?\:\-\(\)]', '', text)
    text = re.sub(r':\s*$', '', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


def number_headings(headings: list[tuple[int, str]]) -> list[tuple[int, str, str]]:
    cleaned = [(lvl, clean_heading_text(text)) for lvl, text in headings]

    level1_count = sum(1 for lvl, _ in cleaned if lvl == 1)
    has_h3 = any(lvl >= 3 for lvl, _ in cleaned)

    numbered: list[tuple[int, str, str]] = []

    if level1_count == 1 and has_h3:
        counters = [0, 0, 0, 0]
        for lvl, text in cleaned:
            if lvl == 1:
                continue
            new_lvl = lvl - 1
            counters[new_lvl - 1] += 1
            for i in range(new_lvl, 4):
                counters[i] = 0
            num = ".".join(str(counters[i]) for i in range(new_lvl))
            numbered.append((new_lvl, num, text))
        return numbered

    if level1_count == 1:
        subs = [text for lvl, text in cleaned if lvl == 2]
        if not subs:
            top = next(text for lvl, text in cleaned if lvl == 1)
            return [(1, "1", top)]
        return [(1, str(i), text) for i, text in enumerate(subs, 1)]

    counters = [0, 0, 0, 0]
    for lvl, text in cleaned:
        if lvl > 4:
            continue
        counters[lvl - 1] += 1
        for i in range(lvl, 4):
            counters[i] = 0
        num = ".".join(str(counters[i]) for i in range(lvl))
        numbered.append((lvl, num, text))

    return numbered


def write_numbered_docx(original: Path, out_path: Path, numbered):
    doc = Document(original)
    new_doc = Document()

    queues = {}
    for lvl, num, text in numbered:
        queues.setdefault(lvl, []).append((num, text))

    for p in doc.paragraphs:
        name = getattr(p.style, "name", "")
        if name.startswith("Heading"):
            lvl = int(name.split()[-1])
            if lvl in queues and queues[lvl]:
                num, text = queues[lvl].pop(0)
                indent = "   " * (lvl - 1)
                new_doc.add_paragraph(f"{indent}{num}-{text}", style=name)
                continue
        new_doc.add_paragraph(p.text, style=name if name else None)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    new_doc.save(out_path)


def main(argv=None):
    parser = argparse.ArgumentParser()
    parser.add_argument("--src", required=True)
    parser.add_argument("--out", default="output")
    parser.add_argument("--write-docx", action="store_true")
    args = parser.parse_args(argv)

    src = Path(args.src)
    out = Path(args.out)
    out.mkdir(exist_ok=True)

    headings = extract_headings_from_docx(src)
    numbered = number_headings(headings)

    if args.write_docx:
        write_numbered_docx(src, out / f"{src.stem}_numbered.docx", numbered)
